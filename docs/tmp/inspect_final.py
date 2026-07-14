from pathlib import Path
from zipfile import ZipFile
from lxml import etree
from docx import Document

p = Path('output/Monographie_Smart_Faculty.docx')
doc = Document(p)
text = '\n'.join(x.text for x in doc.paragraphs)
headings = [(x.style.name, x.text) for x in doc.paragraphs if x.style.name.startswith('Heading')]
with ZipFile(p) as z:
    bad_xml=[]
    for n in z.namelist():
        if n.endswith('.xml') or n.endswith('.rels'):
            try: etree.fromstring(z.read(n))
            except Exception as e: bad_xml.append((n,str(e)))
    images=[n for n in z.namelist() if n.startswith('word/media/')]
print({
    'bytes': p.stat().st_size,
    'paragraphs': len(doc.paragraphs),
    'characters': len(text),
    'tables': len(doc.tables),
    'sections': len(doc.sections),
    'images': len(images),
    'headings': len(headings),
    'heading_1': sum(s=='Heading 1' for s,_ in headings),
    'heading_2': sum(s=='Heading 2' for s,_ in headings),
    'heading_3': sum(s=='Heading 3' for s,_ in headings),
    'xml_errors': bad_xml,
    'has_summary': 'Cette monographie présente' in text,
    'has_five_chapters': all(f'{i}. ' in text for i in range(1,6)),
    'has_bibliography': '6. Bibliographie et webographie' in text,
    'has_annexes': '7. Annexes' in text,
})
