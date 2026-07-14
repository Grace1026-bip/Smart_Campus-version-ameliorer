from pathlib import Path
from datetime import date
from math import ceil
from PIL import Image, ImageDraw, ImageFont

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TMP = ROOT / "tmp" / "monograph"
OUT = ROOT / "output"
TMP.mkdir(parents=True, exist_ok=True)
OUT.mkdir(parents=True, exist_ok=True)
FINAL = OUT / "Monographie_Smart_Faculty.docx"

NAVY = "0B4A7A"
CYAN = "178CA4"
GOLD = "C79A3B"
INK = "1D2A35"
MUTED = "5F6B76"
LIGHT = "F3F6F8"
PALE_BLUE = "EAF2F7"
PALE_GOLD = "F8F1E3"
WHITE = "FFFFFF"
RED = "A33A3A"
GREEN = "2D7A5B"


def font_path(bold=False):
    base = Path("C:/Windows/Fonts")
    choices = ["arialbd.ttf" if bold else "arial.ttf", "segoeuib.ttf" if bold else "segoeui.ttf"]
    for name in choices:
        p = base / name
        if p.exists():
            return str(p)
    return None


def fnt(size, bold=False):
    path = font_path(bold)
    return ImageFont.truetype(path, size) if path else ImageFont.load_default()


def rounded(draw, box, radius=24, fill=WHITE, outline=None, width=2):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def text_center(draw, box, text, font, fill=INK, spacing=6):
    x1, y1, x2, y2 = box
    lines = text.split("\n")
    heights = []
    widths = []
    for line in lines:
        b = draw.textbbox((0, 0), line, font=font)
        widths.append(b[2] - b[0])
        heights.append(b[3] - b[1])
    total = sum(heights) + spacing * (len(lines) - 1)
    y = y1 + (y2 - y1 - total) / 2
    for line, w, h in zip(lines, widths, heights):
        draw.text((x1 + (x2 - x1 - w) / 2, y), line, font=font, fill=fill)
        y += h + spacing


def arrow(draw, start, end, color=NAVY, width=7):
    draw.line([start, end], fill="#" + color, width=width)
    x, y = end
    if abs(end[0] - start[0]) > abs(end[1] - start[1]):
        s = -1 if end[0] > start[0] else 1
        draw.polygon([(x, y), (x + 18*s, y-12), (x + 18*s, y+12)], fill="#" + color)
    else:
        s = -1 if end[1] > start[1] else 1
        draw.polygon([(x, y), (x-12, y + 18*s), (x+12, y + 18*s)], fill="#" + color)


def make_architecture(path):
    im = Image.new("RGB", (1800, 1000), "white")
    d = ImageDraw.Draw(im)
    d.text((90, 55), "Architecture logique de Smart Faculty", font=fnt(46, True), fill="#" + NAVY)
    d.text((90, 115), "Séparation de l'interface, de la logique métier et des données", font=fnt(25), fill="#" + MUTED)
    boxes = [
        (120, 270, 520, 720, "APPLICATION FLUTTER", "Interfaces par rôle\nNavigation et états\nServices d'accès à l'interface"),
        (700, 270, 1100, 720, "API FASTAPI", "Authentification\nRègles métier\nCalculs académiques\nContrôle des permissions"),
        (1280, 270, 1680, 720, "BASE MYSQL", "Données relationnelles\nContraintes d'intégrité\nMigrations Alembic\nTransactions InnoDB"),
    ]
    fills = [PALE_BLUE, "EEF7F5", PALE_GOLD]
    accents = [NAVY, CYAN, GOLD]
    for (x1,y1,x2,y2,title,body), fill, accent in zip(boxes, fills, accents):
        rounded(d, (x1,y1,x2,y2), 30, "#"+fill, "#"+accent, 4)
        d.rectangle((x1,y1,x2,y1+16), fill="#"+accent)
        text_center(d, (x1+25,y1+55,x2-25,y1+155), title, fnt(31, True), "#"+accent)
        text_center(d, (x1+35,y1+170,x2-35,y2-35), body, fnt(28), "#"+INK, 16)
    arrow(d, (540, 495), (675, 495), NAVY)
    text_center(d, (525, 380, 690, 460), "HTTP / JSON", fnt(21, True), "#"+MUTED)
    arrow(d, (1120, 495), (1255, 495), CYAN)
    text_center(d, (1110, 380, 1265, 460), "SQL", fnt(21, True), "#"+MUTED)
    rounded(d, (410, 820, 1390, 925), 22, "#"+NAVY, None, 0)
    text_center(d, (430, 830, 1370, 915), "JWT + contrôle du rôle + journalisation + tests automatisés", fnt(27, True), "white")
    im.save(path, quality=95)


def make_roles(path):
    im = Image.new("RGB", (1800, 1100), "white")
    d = ImageDraw.Draw(im)
    d.text((90, 55), "Acteurs et responsabilités", font=fnt(46, True), fill="#" + NAVY)
    center = (700, 420, 1100, 680)
    rounded(d, center, 32, "#"+NAVY, None, 0)
    text_center(d, center, "SMART\nFACULTY", fnt(38, True), "white", 10)
    roles = [
        ((100,220,500,430), "ADMINISTRATEUR", "Comptes, rôles, paramètres", NAVY),
        ((1300,220,1700,430), "DOYEN", "Indicateurs et supervision", GOLD),
        ((100,760,500,970), "APPARITEUR", "Enrôlements et opérations", CYAN),
        ((1300,760,1700,970), "ENSEIGNANT", "Cours, valve, notes, projets", GREEN),
        ((700,820,1100,1030), "ÉTUDIANT", "Résultats, alertes, réclamations", RED),
    ]
    for box,title,body,color in roles:
        rounded(d, box, 25, "#"+LIGHT, "#"+color, 4)
        text_center(d, (box[0]+20,box[1]+25,box[2]-20,box[1]+100), title, fnt(27, True), "#"+color)
        text_center(d, (box[0]+25,box[1]+105,box[2]-25,box[3]-20), body, fnt(23), "#"+INK)
        sx = (box[0]+box[2])//2
        sy = box[3] if box[1] < 500 else box[1]
        ex = (center[0]+center[2])//2 if abs(sx-900)<120 else (center[0] if sx<900 else center[2])
        ey = center[1] if sy<center[1] else center[3]
        arrow(d, (sx,sy), (ex,ey), color, 5)
    im.save(path, quality=95)


def make_roadmap(path):
    im = Image.new("RGB", (1800, 800), "white")
    d = ImageDraw.Draw(im)
    d.text((90, 55), "Démarche de réalisation", font=fnt(46, True), fill="#"+NAVY)
    stages = [
        ("1", "AUDIT", "Cartographier\nl'existant"),
        ("2", "CADRAGE", "Stabiliser\nl'architecture"),
        ("3", "CONCEPTION", "Définir données\net règles"),
        ("4", "IMPLÉMENTATION", "Développer par\nmodules"),
        ("5", "VALIDATION", "Tester, rendre\net contrôler"),
    ]
    y = 390
    d.line((170,y,1630,y), fill="#"+PALE_BLUE, width=24)
    for i,(n,title,body) in enumerate(stages):
        x = 170 + i*365
        d.ellipse((x-50,y-50,x+50,y+50), fill="#"+NAVY if i<4 else "#"+GOLD)
        text_center(d, (x-45,y-45,x+45,y+45), n, fnt(32, True), "white")
        text_center(d, (x-150,200,x+150,310), title, fnt(25, True), "#"+(NAVY if i<4 else GOLD))
        text_center(d, (x-155,465,x+155,610), body, fnt(24), "#"+INK, 8)
    rounded(d, (260,665,1540,750), 18, "#"+LIGHT, None, 0)
    text_center(d, (280,675,1520,740), "Chaque module est documenté, migré sur la base de test puis vérifié avant intégration.", fnt(24, True), "#"+MUTED)
    im.save(path, quality=95)


def make_tests(path):
    im = Image.new("RGB", (1800, 1000), "white")
    d = ImageDraw.Draw(im)
    d.text((90, 45), "Progression des suites automatisées", font=fnt(46, True), fill="#"+NAVY)
    d.text((90, 105), "Résultats consignés dans le journal de développement", font=fnt(25), fill="#"+MUTED)
    labels = ["Socle", "Auth.", "Inscr.", "CORS", "Ens.", "Notes", "Sem.", "LMD", "Projet", "Enrôl.", "Encadr.", "Étudiant"]
    backend = [26, 41, 57, 61, 68, 73, 97, 107, 120, 128, 134, 141]
    flutter = [2, 12, 15, 24, 28, 35, 36, 37, 39, 42, 44, 47]
    x0,y0,x1,y1 = 150,210,1690,810
    d.line((x0,y0,x0,y1), fill="#"+INK, width=3)
    d.line((x0,y1,x1,y1), fill="#"+INK, width=3)
    maxv=150
    for v in range(0,151,25):
        y=y1-(y1-y0)*v/maxv
        d.line((x0,y,x1,y), fill="#D9E2E8", width=2)
        d.text((75,y-16), str(v), font=fnt(20), fill="#"+MUTED)
    xs=[]
    for i,lbl in enumerate(labels):
        x=x0+(x1-x0)*i/(len(labels)-1)
        xs.append(x)
        d.text((x-38,y1+22), lbl, font=fnt(18), fill="#"+MUTED)
    def plot(vals,color):
        pts=[]
        for x,v in zip(xs,vals):
            y=y1-(y1-y0)*v/maxv
            pts.append((x,y))
        d.line(pts, fill="#"+color, width=7)
        for (x,y),v in zip(pts,vals):
            d.ellipse((x-9,y-9,x+9,y+9), fill="white", outline="#"+color, width=5)
            d.text((x-15,y-35), str(v), font=fnt(16, True), fill="#"+color)
    plot(backend,NAVY); plot(flutter,GOLD)
    rounded(d,(1110,95,1380,165),14,"#"+PALE_BLUE,None,0); text_center(d,(1120,100,1370,160),"Backend",fnt(22,True),"#"+NAVY)
    rounded(d,(1400,95,1670,165),14,"#"+PALE_GOLD,None,0); text_center(d,(1410,100,1660,160),"Flutter",fnt(22,True),"#"+GOLD)
    d.text((150,885), "Dernier état validé : 141 tests backend et 47 tests Flutter, deux exécutions complètes.", font=fnt(24, True), fill="#"+INK)
    im.save(path, quality=95)


def make_lmd(path):
    im = Image.new("RGB", (1800, 720), "white")
    d = ImageDraw.Draw(im)
    d.text((90, 45), "Chaîne de décision académique", font=fnt(46, True), fill="#"+NAVY)
    stages = [
        ("ENSEIGNANT", "Publie et verrouille\nles résultats", NAVY),
        ("MOTEUR", "Calcule moyenne\net crédits", CYAN),
        ("JURY", "Décide : admission,\ncompensation ou ajournement", GOLD),
        ("APPARITEUR", "Publie la session\nclôturée", GREEN),
        ("ÉTUDIANT", "Consulte le snapshot\nofficiel", RED),
    ]
    for i,(title,body,color) in enumerate(stages):
        x=70+i*350
        box=(x,230,x+290,500)
        rounded(d,box,24,"#"+LIGHT,"#"+color,4)
        text_center(d,(x+15,250,x+275,325),title,fnt(24,True),"#"+color)
        text_center(d,(x+20,335,x+270,480),body,fnt(22),"#"+INK,8)
        if i<4: arrow(d,(x+300,365),(x+340,365),color,5)
    rounded(d,(250,590,1550,665),16,"#"+PALE_BLUE,None,0)
    text_center(d,(270,595,1530,660),"La clôture produit une version immuable ; toute correction crée une nouvelle version historisée.",fnt(23,True),"#"+NAVY)
    im.save(path, quality=95)


ARCH = TMP / "architecture.png"
ROLES = TMP / "roles.png"
ROADMAP = TMP / "roadmap.png"
TESTS = TMP / "tests.png"
LMD = TMP / "lmd_workflow.png"
make_architecture(ARCH)
make_roles(ROLES)
make_roadmap(ROADMAP)
make_tests(TESTS)
make_lmd(LMD)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=110, start=120, bottom=110, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v)); node.set(qn("w:type"), "dxa")


def repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def set_table_geometry(table, widths_dxa, indent=120):
    total = sum(widths_dxa)
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW"); tblPr.append(tblW)
    tblW.set(qn("w:w"), str(total)); tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd"); tblPr.append(tblInd)
    tblInd.set(qn("w:w"), str(indent)); tblInd.set(qn("w:type"), "dxa")
    layout = tblPr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout"); tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for w in widths_dxa:
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(w)); grid.append(col)
    for row in table.rows:
        for cell,w in zip(row.cells,widths_dxa):
            tcPr=cell._tc.get_or_add_tcPr(); tcW=tcPr.find(qn("w:tcW"))
            if tcW is None: tcW=OxmlElement("w:tcW"); tcPr.append(tcW)
            tcW.set(qn("w:w"),str(w)); tcW.set(qn("w:type"),"dxa")
            set_cell_margins(cell)


def set_font(run, name="Times New Roman", size=12, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic


def set_repeat_together(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement("w:keepNext")
    pPr.append(keep)


def add_field(paragraph, instruction, placeholder=""):
    r = paragraph.add_run()
    fldChar = OxmlElement("w:fldChar"); fldChar.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText"); instrText.set(qn("xml:space"), "preserve"); instrText.text = instruction
    separate = OxmlElement("w:fldChar"); separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t"); text.text = placeholder
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    for node in (fldChar, instrText, separate, text, end): r._r.append(node)


def set_page_num(section, fmt="decimal", start=1):
    sectPr = section._sectPr
    pg = sectPr.find(qn("w:pgNumType"))
    if pg is None:
        pg = OxmlElement("w:pgNumType"); sectPr.append(pg)
    pg.set(qn("w:fmt"), fmt); pg.set(qn("w:start"), str(start))


def set_a4(section):
    section.page_width = Cm(21.0); section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5); section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0); section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.25); section.footer_distance = Cm(1.25)


doc = Document()
for s in doc.sections: set_a4(s)

# Document metadata
doc.core_properties.title = "Conception et développement de Smart Faculty"
doc.core_properties.subject = "Monographie de Licence 2 - Sciences Informatiques"
doc.core_properties.author = "Étudiant(e) - à compléter"
doc.core_properties.keywords = "Smart Faculty, gestion académique, Flutter, FastAPI, MySQL, LMD"
doc.core_properties.comments = "Monographie construite à partir du corpus documentaire du projet Smart Faculty."

# Styles: narrative_proposal with named academic_monograph overrides.
styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Times New Roman"; normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman"); normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
normal.font.size = Pt(12); normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
normal.paragraph_format.space_before = Pt(0); normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.5
normal.paragraph_format.widow_control = True

for name,size,before,after,color in [
    ("Title",26,0,8,NAVY),("Subtitle",14,0,8,MUTED),
    ("Heading 1",16,18,10,NAVY),("Heading 2",13,12,6,NAVY),("Heading 3",12,8,4,"1F4D78")]:
    st=styles[name]; st.font.name="Arial"; st._element.rPr.rFonts.set(qn("w:ascii"),"Arial"); st._element.rPr.rFonts.set(qn("w:hAnsi"),"Arial")
    st.font.size=Pt(size); st.font.bold=True if name!="Subtitle" else False; st.font.color.rgb=RGBColor.from_string(color)
    st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after)
    st.paragraph_format.keep_with_next=True
styles["Heading 1"].paragraph_format.page_break_before = True
styles["Heading 1"].paragraph_format.outline_level = 0
styles["Heading 2"].paragraph_format.outline_level = 1
styles["Heading 3"].paragraph_format.outline_level = 2

caption = styles["Caption"]
caption.font.name="Times New Roman"; caption._element.rPr.rFonts.set(qn("w:ascii"),"Times New Roman"); caption._element.rPr.rFonts.set(qn("w:hAnsi"),"Times New Roman")
caption.font.size=Pt(10); caption.font.italic=True; caption.font.color.rgb=RGBColor.from_string(MUTED)
caption.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER; caption.paragraph_format.space_before=Pt(4); caption.paragraph_format.space_after=Pt(10)

if "Résumé" not in styles:
    st = styles.add_style("Résumé", WD_STYLE_TYPE.PARAGRAPH)
else: st=styles["Résumé"]
st.font.name="Times New Roman"; st._element.rPr.rFonts.set(qn("w:ascii"),"Times New Roman"); st._element.rPr.rFonts.set(qn("w:hAnsi"),"Times New Roman")
st.font.size=Pt(12); st.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY; st.paragraph_format.line_spacing=1.5; st.paragraph_format.space_after=Pt(8)

for list_name in ("List Bullet", "List Number"):
    st=styles[list_name]; st.font.name="Times New Roman"; st._element.rPr.rFonts.set(qn("w:ascii"),"Times New Roman"); st._element.rPr.rFonts.set(qn("w:hAnsi"),"Times New Roman")
    st.font.size=Pt(12); st.paragraph_format.left_indent=Cm(0.95); st.paragraph_format.first_line_indent=Cm(-0.48); st.paragraph_format.space_after=Pt(4); st.paragraph_format.line_spacing=1.25


def p(text="", style=None, align=None, bold_lead=None, italic=False):
    par = doc.add_paragraph(style=style)
    if align is not None: par.alignment=align
    if bold_lead and text.startswith(bold_lead):
        r=par.add_run(bold_lead); set_font(r,bold=True)
        r=par.add_run(text[len(bold_lead):]); set_font(r,italic=italic)
    else:
        r=par.add_run(text); set_font(r,italic=italic)
    return par


def heading(text, level=1):
    return doc.add_heading(text, level=level)


def bullets(items):
    for item in items: p(item, "List Bullet")


def numbers(items):
    for item in items: p(item, "List Number")


def callout(title, text, fill=PALE_BLUE, accent=NAVY):
    par=doc.add_paragraph()
    par.paragraph_format.space_before=Pt(6); par.paragraph_format.space_after=Pt(10); par.paragraph_format.left_indent=Cm(0.35); par.paragraph_format.right_indent=Cm(0.25)
    pPr=par._p.get_or_add_pPr(); shd=OxmlElement("w:shd"); shd.set(qn("w:fill"),fill); pPr.append(shd)
    borders=OxmlElement("w:pBdr"); left=OxmlElement("w:left"); left.set(qn("w:val"),"single"); left.set(qn("w:sz"),"22"); left.set(qn("w:color"),accent); borders.append(left); pPr.append(borders)
    r=par.add_run(title+" "); set_font(r,"Arial",11,accent,True)
    r=par.add_run(text); set_font(r,"Times New Roman",11,INK)
    return par


def table(headers, rows, widths, aligns=None, font_size=10.5, caption_text=None):
    t=doc.add_table(rows=1, cols=len(headers))
    t.alignment=WD_TABLE_ALIGNMENT.LEFT; t.autofit=False
    t.style="Table Grid"
    hdr=t.rows[0]; repeat_table_header(hdr)
    for j,h in enumerate(headers):
        c=hdr.cells[j]; c.text=""; set_cell_shading(c,NAVY); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        pr=c.paragraphs[0]; pr.alignment=WD_ALIGN_PARAGRAPH.CENTER; pr.paragraph_format.space_after=Pt(0); pr.paragraph_format.line_spacing=1.0
        rr=pr.add_run(h); set_font(rr,"Arial",10,WHITE,True)
    for i,row in enumerate(rows):
        cells=t.add_row().cells
        for j,val in enumerate(row):
            c=cells[j]; c.text=""; c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if i%2: set_cell_shading(c,"F8FAFB")
            pr=c.paragraphs[0]; pr.paragraph_format.space_after=Pt(0); pr.paragraph_format.line_spacing=1.05
            pr.alignment=(aligns[j] if aligns else WD_ALIGN_PARAGRAPH.LEFT)
            rr=pr.add_run(str(val)); set_font(rr,"Times New Roman",font_size,INK)
    set_table_geometry(t,widths)
    if caption_text:
        doc.add_paragraph(caption_text, style="Caption")
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
    return t


def figure(path, caption_text, width_cm=15.3, alt=""):
    par=doc.add_paragraph(); par.alignment=WD_ALIGN_PARAGRAPH.CENTER; par.paragraph_format.space_before=Pt(6); par.paragraph_format.space_after=Pt(2); set_repeat_together(par)
    run=par.add_run(); inline=run.add_picture(str(path), width=Cm(width_cm))
    docPr=inline._inline.docPr
    if alt: docPr.set("descr", alt)
    cp=doc.add_paragraph(caption_text, style="Caption")
    return cp


def equation(text):
    par=p(text,align=WD_ALIGN_PARAGRAPH.CENTER)
    for r in par.runs: set_font(r,"Cambria Math",11,NAVY,True)
    par.paragraph_format.space_before=Pt(5); par.paragraph_format.space_after=Pt(8)


def section_header_footer(section, label, fmt="decimal", start=1):
    section.header.is_linked_to_previous=False; section.footer.is_linked_to_previous=False
    hp=section.header.paragraphs[0]; hp.text=""; hp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    r=hp.add_run(label.upper()); set_font(r,"Arial",8.5,MUTED,True)
    hp.paragraph_format.space_after=Pt(0)
    fp=section.footer.paragraphs[0]; fp.text=""; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=fp.add_run("SMART FACULTY  |  "); set_font(r,"Arial",8.5,MUTED)
    add_field(fp,"PAGE","1")
    set_page_num(section,fmt,start)


# Cover section
cover=doc.sections[0]
cover.header.is_linked_to_previous=False; cover.footer.is_linked_to_previous=False
cover.header.paragraphs[0].text=""; cover.footer.paragraphs[0].text=""
p("UNIVERSITÉ PROTESTANTE AU CONGO",align=WD_ALIGN_PARAGRAPH.CENTER).paragraph_format.space_before=Pt(6)
q=p("FACULTÉ DES SCIENCES INFORMATIQUES",align=WD_ALIGN_PARAGRAPH.CENTER); q.paragraph_format.space_after=Pt(4)
for r in q.runs: set_font(r,"Arial",11,NAVY,True)
q=p("LICENCE 2 - SYSTÈME LMD",align=WD_ALIGN_PARAGRAPH.CENTER); q.paragraph_format.space_after=Pt(50)
for r in q.runs: set_font(r,"Arial",10,MUTED,True)

q=p("MONOGRAPHIE",align=WD_ALIGN_PARAGRAPH.CENTER); q.paragraph_format.space_after=Pt(12)
for r in q.runs: set_font(r,"Arial",13,GOLD,True)
q=p("CONCEPTION ET DÉVELOPPEMENT\nDE SMART FACULTY",align=WD_ALIGN_PARAGRAPH.CENTER); q.paragraph_format.space_after=Pt(12)
for r in q.runs: set_font(r,"Arial",25,NAVY,True)
q=p("Une plateforme intelligente de gestion académique\npour une faculté universitaire",align=WD_ALIGN_PARAGRAPH.CENTER); q.paragraph_format.space_after=Pt(48)
for r in q.runs: set_font(r,"Times New Roman",15,INK,False,True)

callout("OBJET DU TRAVAIL", "Centraliser les services académiques, sécuriser les opérations par rôle et soutenir la décision grâce à une architecture Flutter, FastAPI et MySQL.", PALE_BLUE, NAVY)
p("").paragraph_format.space_after=Pt(20)
meta = table(["PRÉSENTÉE PAR", "SOUS LA DIRECTION DE"], [
    ("Nom et prénom : __________________________\nMatricule : _______________________________", "Encadreur : ______________________________\nQualité : _________________________________"),
], [4560,4560], font_size=10.5)
p("").paragraph_format.space_after=Pt(20)
q=p("Année académique 2025-2026",align=WD_ALIGN_PARAGRAPH.CENTER)
for r in q.runs: set_font(r,"Arial",11,NAVY,True)

# Front matter section
front = doc.add_section(WD_SECTION.NEW_PAGE); set_a4(front); section_header_footer(front,"Monographie - Smart Faculty","lowerRoman",1)
front.different_first_page_header_footer=False
styles["Heading 1"].paragraph_format.page_break_before=False
heading("Résumé",1)
p("Cette monographie présente la conception et la réalisation de Smart Faculty, une plateforme de gestion académique destinée à centraliser les services d'une faculté universitaire. Le travail s'appuie sur une analyse des besoins des étudiants, des enseignants et des responsables académiques, puis sur une architecture séparant l'interface utilisateur, le serveur applicatif et la base de données. La solution obtenue sécurise les accès selon les responsabilités, automatise plusieurs traitements académiques, facilite la publication des résultats et organise les enrôlements ainsi que les encadrements de projets. Une démarche progressive de développement, de migration et de vérification a permis de stabiliser les modules sans compromettre les données existantes. Les validations consignées montrent un système cohérent, testable et évolutif, tout en faisant apparaître des limites relatives au déploiement en production et à certaines fonctions avancées encore à réaliser.","Résumé")
q=p("Mots-clés : gestion académique, transformation numérique, application web, sécurité, résultats académiques, faculté universitaire.")
for r in q.runs: set_font(r,"Times New Roman",11,MUTED,False,True)

heading("Table des matières",1)
toc=doc.add_paragraph(); add_field(toc,'TOC \\o "1-3" \\h \\z \\u',"La table des matières sera actualisée à l'ouverture du document.")
toc.paragraph_format.space_after=Pt(12)
callout("UTILISATION DANS WORD", "Si les numéros de page ne se mettent pas à jour automatiquement, sélectionner la table puis choisir « Mettre à jour toute la table ».", PALE_GOLD, GOLD)

heading("Liste des figures",1)
figs=[
    "Figure 1. Acteurs et responsabilités de Smart Faculty",
    "Figure 2. Démarche de réalisation",
    "Figure 3. Architecture logique de Smart Faculty",
    "Figure 4. Chaîne de décision académique",
    "Figure 5. Progression des suites automatisées",
]
bullets(figs)
heading("Liste des tableaux",1)
bullets([
    "Tableau 1. Rôles et responsabilités principales",
    "Tableau 2. Synthèse des besoins fonctionnels",
    "Tableau 3. Exigences non fonctionnelles",
    "Tableau 4. Technologies et rôles dans la solution",
    "Tableau 5. Domaines de données principaux",
    "Tableau 6. Principales étapes de réalisation",
    "Tableau 7. Résultats de validation",
    "Tableau 8. Difficultés, réponses et enseignements",
    "Tableau 9. Limites actuelles et évolutions associées",
])

heading("Sigles et abréviations",1)
table(["Sigle", "Signification"], [
    ("API", "Interface de programmation d'application"),
    ("HTTP", "Protocole de transfert hypertexte"),
    ("JSON", "Notation textuelle d'échange de données"),
    ("JWT", "Jeton web signé utilisé pour l'authentification"),
    ("LMD", "Licence - Maîtrise - Doctorat"),
    ("MVP", "Version minimale viable d'un produit"),
    ("REST", "Style d'architecture pour services web"),
    ("SQL", "Langage de requête structuré"),
    ("UE", "Unité d'enseignement"),
], [1800,7320])

# Body section
body = doc.add_section(WD_SECTION.NEW_PAGE); set_a4(body); section_header_footer(body,"Smart Faculty - Monographie L2","decimal",1)
styles["Heading 1"].paragraph_format.page_break_before=True

heading("1. Introduction générale",1)
heading("1.1 Contexte",2)
p("La transformation numérique des établissements d'enseignement supérieur répond à un besoin concret : mieux organiser l'information, réduire les opérations manuelles et rendre les services académiques plus accessibles. Dans une faculté, les données relatives aux étudiants, aux cours, aux résultats, aux annonces, aux réclamations et aux décisions de jury sont produites par plusieurs acteurs. Lorsqu'elles restent dispersées entre documents, messages et outils non intégrés, leur traitement devient lent, difficile à contrôler et peu favorable à la prise de décision.")
p("Le projet Smart Faculty s'inscrit dans ce contexte. Il vise à réunir les principales activités académiques dans une plateforme unique, adaptée aux étudiants, enseignants, appariteurs, doyens et administrateurs. La vision formulée dans les documents d'admission consiste à proposer une solution institutionnelle, sécurisée, intuitive et capable d'évoluer progressivement avec les besoins de la faculté (Équipe Smart Faculty, 2026a). Cette ambition rejoint la finalité pédagogique du guide de monographie : mobiliser les compétences acquises pour résoudre un problème réel par une démarche structurée (Faculté des Sciences Informatiques, 2025).")
p("Le travail ne se limite pas à la réalisation d'interfaces. Il comprend la clarification d'une architecture existante, la consolidation de règles académiques, la protection des données, l'isolation des tests et l'implémentation progressive de modules cohérents. Le journal de développement retrace ainsi le passage d'un dépôt comportant plusieurs technologies concurrentes à une architecture officielle fondée sur Flutter, FastAPI et MySQL (Équipe Smart Faculty, 2026b).")
figure(ROLES,"Figure 1. Acteurs et responsabilités de Smart Faculty. Source : élaboration à partir des documents d'analyse du projet.",15.2,"Schéma des cinq acteurs de Smart Faculty autour de la plateforme")

heading("1.2 Problématique",2)
p("Dans de nombreuses structures universitaires, la gestion académique demeure fragmentée. Les enseignants publient les informations selon des canaux variés, les étudiants disposent de peu de visibilité sur leurs résultats et leurs réclamations, tandis que les responsables doivent consolider manuellement des données avant de décider. Cette fragmentation augmente les risques d'erreur, de retard, d'accès non autorisé et de perte de traçabilité.")
p("La problématique centrale de cette monographie est donc formulée comme suit : comment concevoir et réaliser une plateforme intelligente capable de centraliser les activités académiques d'une faculté, de sécuriser les opérations selon les responsabilités de chaque acteur, d'automatiser les calculs et workflows essentiels, tout en restant maintenable, testable et évolutive ?")
callout("QUESTION DIRECTRICE", "Comment transformer un ensemble de besoins académiques hétérogènes en un système cohérent qui protège l'autorité, les données et l'historique des décisions ?", PALE_GOLD, GOLD)

heading("1.3 Objectifs",2)
heading("1.3.1 Objectif général",3)
p("Concevoir et développer une plateforme de gestion académique permettant de centraliser les informations d'une faculté universitaire, de faciliter les activités des utilisateurs et de soutenir la prise de décision grâce à des traitements sécurisés et vérifiables.")
heading("1.3.2 Objectifs spécifiques",3)
bullets([
    "mettre en place une authentification sécurisée et une autorisation fondée sur le rôle actif de l'utilisateur ;",
    "organiser les comptes, promotions, cours, inscriptions et enrôlements académiques ;",
    "permettre aux enseignants de gérer leurs cours, publications, évaluations et résultats ;",
    "offrir aux étudiants un accès contrôlé à leurs informations, résultats, projets et documents officiels ;",
    "automatiser les moyennes, crédits et étapes de délibération selon les règles académiques retenues ;",
    "faciliter le traitement des réclamations, notifications et alertes académiques ;",
    "garantir la cohérence des données au moyen d'une base relationnelle et de migrations contrôlées ;",
    "valider chaque évolution par des tests automatisés, des vérifications de migration et des contrôles visuels."
])

heading("1.4 Méthodologie adoptée",2)
p("La démarche suivie combine analyse documentaire, conception fonctionnelle, développement incrémental et validation continue. Les documents d'admission ont d'abord permis de préciser la vision, le périmètre et les utilisateurs. Les documents d'analyse ont ensuite défini les besoins, les cas d'utilisation, les scénarios et les règles métier. Les documents de conception ont fourni la structure client-serveur, l'organisation du projet, la base de données et les interfaces de programmation.")
p("La réalisation s'est déroulée par interventions ciblées. Chaque module a été audité avant modification ; les comportements existants jugés corrects ont été conservés ; les ajouts de données ont pris la forme de migrations additives ; les tests ont été exécutés exclusivement sur une base portant le suffixe de test. Cette discipline a évité de mélanger diagnostic et correction, et a permis de mesurer la progression du système à chaque étape.")
figure(ROADMAP,"Figure 2. Démarche de réalisation. Source : synthèse du journal de développement.",15.3,"Frise en cinq étapes : audit, cadrage, conception, implémentation et validation")

heading("1.5 Structure du document",2)
p("Outre la présente introduction, le document comprend quatre chapitres. Le deuxième présente les concepts et les technologies mobilisés. Le troisième expose l'analyse des besoins, l'architecture, les étapes d'implémentation et les résultats obtenus. Le quatrième examine les difficultés rencontrées et les limites de la version actuelle. Le cinquième récapitule les apports du projet et propose des perspectives. Une bibliographie, une webographie et des annexes complètent le travail conformément au guide de rédaction.")

heading("2. Cadre théorique et technologique",1)
heading("2.1 Définitions des concepts clés",2)
heading("2.1.1 Système d'information académique",3)
p("Un système d'information académique regroupe les personnes, procédures, données et moyens logiciels nécessaires à la gestion des activités pédagogiques et administratives. Dans Smart Faculty, cette notion recouvre l'identité des utilisateurs, les promotions, les cours, les évaluations, les résultats, les publications, les réclamations, les enrôlements et les projets. La valeur du système tient moins à l'accumulation de fonctions qu'à la cohérence des échanges entre ces domaines.")

heading("2.1.2 Architecture client-serveur",3)
p("L'architecture client-serveur sépare l'interface utilisée par la personne du service chargé d'appliquer les règles et de manipuler les données. Le client Flutter présente les écrans et transmet les actions. Le serveur FastAPI authentifie l'utilisateur, contrôle ses permissions, applique la logique métier et accède à MySQL. Cette séparation évite que la sécurité dépende uniquement de l'interface et facilite l'évolution indépendante des composants.")

heading("2.1.3 Interface REST et échange JSON",3)
p("Une interface de programmation REST organise les opérations sous forme de ressources accessibles par des méthodes HTTP. Smart Faculty regroupe les routes sous le préfixe /api/v1 et échange des données structurées au format JSON. FastAPI s'appuie sur des standards ouverts, notamment OpenAPI et JSON Schema, et produit une documentation interactive utile à la vérification des routes (FastAPI, 2026).")

heading("2.1.4 Gestion des rôles et principe de moindre privilège",3)
p("Le contrôle d'accès fondé sur les rôles associe chaque action à une responsabilité vérifiée. Dans Smart Faculty, l'interface peut demander un rôle actif, mais le backend recharge les rôles depuis la base et refuse tout rôle non possédé. L'étudiant ne consulte que ses propres données ; l'enseignant agit sur ses cours ; l'appariteur exécute les opérations administratives ; le doyen supervise ; l'administrateur assure les fonctions techniques. Ce principe réduit l'exposition des données et empêche le client de devenir une source d'autorité.")

heading("2.1.5 Base de données relationnelle et intégrité",3)
p("Une base de données relationnelle organise les informations dans des tables liées par des clés. Les contraintes d'unicité, les clés étrangères et les transactions protègent la cohérence des opérations. La documentation MySQL rappelle que les clés étrangères relient des données distribuées entre tables et que le moteur transactionnel InnoDB peut annuler une instruction qui viole certaines contraintes (Oracle, 2026). Cette propriété s'est révélée importante pour l'isolation des tests de Smart Faculty.")

heading("2.1.6 Migration de schéma",3)
p("Une migration décrit une transformation versionnée de la structure de la base. Alembic est utilisé pour ajouter progressivement des tables ou colonnes, contrôler l'ordre des évolutions et tester les montées et descentes de version. Le projet privilégie les migrations additives, précédées d'une sauvegarde lorsque la base principale est concernée.")

heading("2.1.7 Système LMD, crédits et délibération",3)
p("Le système Licence - Maîtrise - Doctorat organise les apprentissages en semestres et crédits capitalisables. Pour le moteur actuel, les résultats historiques sur cent sont convertis sur vingt. Le seuil d'acquisition retenu est de dix sur vingt, la moyenne semestrielle est pondérée par les crédits et les décisions finales sont l'admission, l'admission avec compensation, la défaillance ou l'ajournement. Ces règles sont documentées à partir des textes communiqués par la maîtrise d'ouvrage et de l'Instruction académique numéro 027 pour 2025-2026 (Ministère de l'Enseignement supérieur, 2025 ; Équipe Smart Faculty, 2026c).")
equation("Moyenne semestrielle = Σ (note du cours × crédits du cours) / Σ crédits")

heading("2.1.8 Test automatisé et traçabilité",3)
p("Un test automatisé vérifie qu'un comportement attendu reste correct après une modification. Le projet distingue les tests backend, les tests Flutter, l'analyse statique, la compilation web, les contrôles HTTP et les cycles de migration. La traçabilité associe chaque évolution à ses règles, fichiers, risques, résultats et limites. Elle permet de présenter des preuves reproductibles plutôt qu'une simple appréciation visuelle.")

heading("2.2 Présentation des outils ou technologies utilisés",2)
table(["Technologie", "Rôle dans Smart Faculty", "Motif du choix"], [
    ("Flutter / Dart", "Interfaces web et écrans adaptés aux rôles", "Composants réutilisables, multiplateforme et testabilité"),
    ("Python / FastAPI", "Routes, validation, sécurité et logique métier", "API typée, documentation automatique et intégration simple"),
    ("SQLAlchemy", "Modèles et accès relationnel aux données", "Séparation entre logique métier et requêtes"),
    ("Alembic", "Versionnement du schéma", "Migrations contrôlées et réversibles"),
    ("MySQL / InnoDB", "Stockage persistant et transactions", "Relations, contraintes, index et disponibilité locale"),
    ("JWT et bcrypt", "Sessions signées et protection des mots de passe", "Contrôle stateless avec secrets non exposés"),
    ("ReportLab", "Génération de fiches académiques en PDF", "Production en mémoire et contrôle du rendu"),
    ("Pytest / flutter test", "Validation backend et frontend", "Exécution automatisée et reproductible"),
    ("Git", "Historique et sauvegarde du code", "Traçabilité des changements et branches de sécurité"),
], [1900,3500,3720], font_size=9.5, caption_text="Tableau 1. Technologies et rôles dans la solution.")
p("L'architecture Flutter est cohérente avec les recommandations officielles de séparation entre couche d'interface et couche de données. Les vues présentent l'état ; les services et dépôts gèrent l'accès aux données et les traitements associés (Flutter, 2026). Le projet adapte cette recommandation à sa structure existante en français, sans réécriture systématique des modules fonctionnels.")
callout("CHOIX D'ARCHITECTURE", "Flutter est le client officiel, FastAPI le backend actif et MySQL la base relationnelle. Les anciens composants PHP et Flask ont été archivés afin de supprimer l'ambiguïté sans perdre l'historique.", PALE_BLUE, NAVY)

heading("3. Réalisation du projet",1)
heading("3.1 Analyse des besoins",2)
heading("3.1.1 Utilisateurs concernés",3)
p("L'analyse fonctionnelle identifie cinq acteurs principaux. Tous partagent l'authentification, la consultation du profil, les notifications et la déconnexion, mais leurs opérations sensibles restent distinctes. Cette séparation structure les écrans, les routes et les tests d'autorisation.")
table(["Acteur", "Responsabilité principale", "Exemples d'opérations"], [
    ("Administrateur", "Administration technique", "Comptes, rôles, paramètres et journaux"),
    ("Doyen", "Supervision académique", "Indicateurs, délibérations et étudiants à risque"),
    ("Appariteur", "Gestion académique quotidienne", "Promotions, enrôlements, projets et publication administrative"),
    ("Enseignant", "Gestion pédagogique de ses cours", "Valve, évaluations, notes et encadrements"),
    ("Étudiant", "Consultation et interaction personnelle", "Résultats, réclamations, enrôlement et projet"),
], [1600,3000,4520], font_size=10, caption_text="Tableau 2. Rôles et responsabilités principales.")

heading("3.1.2 Besoins fonctionnels",3)
p("Les besoins fonctionnels ont été regroupés en domaines afin de préserver la cohérence entre les documents initiaux et les modules effectivement développés.")
table(["Domaine", "Besoins essentiels", "Bénéficiaires"], [
    ("Identité", "Connexion, sessions, rôles et demandes d'inscription", "Tous les utilisateurs"),
    ("Académique", "Promotions, cours, affectations et enrôlements", "Appariteur, enseignant, étudiant"),
    ("Communication", "Valve, documents et notifications", "Enseignant, étudiant"),
    ("Évaluation", "Évaluations, notes, moyennes, crédits et résultats", "Enseignant, étudiant, responsables"),
    ("Délibération", "Jury, décision, clôture, publication et correction versionnée", "Doyen, jury, appariteur, étudiant"),
    ("Suivi", "Réclamations, présences et risques académiques", "Étudiant et responsables"),
    ("Projets", "Projets, spécialités et attribution des encadreurs", "Appariteur, enseignant, étudiant"),
    ("Pilotage", "Tableaux de bord et statistiques", "Doyen, appariteur, administrateur"),
], [1700,4800,2620], font_size=9.5, caption_text="Tableau 3. Synthèse des besoins fonctionnels.")

heading("3.1.3 Besoins non fonctionnels",3)
table(["Exigence", "Traduction dans le projet"], [
    ("Sécurité", "Mot de passe haché, jetons signés, rôle actif vérifié côté backend et données minimales"),
    ("Fiabilité", "Transactions, contraintes, migrations et historique des décisions"),
    ("Performance", "Requêtes ciblées, chargement relationnel adapté et réponses JSON"),
    ("Ergonomie", "Navigation par rôle, états de chargement, vide, erreur et session expirée"),
    ("Maintenabilité", "Modules séparés, services dédiés, documentation et architecture stable"),
    ("Évolutivité", "Migrations additives, référentiels contrôlés et périmètre MVP explicite"),
    ("Traçabilité", "Journal, snapshots officiels, auteurs, dates et statuts historisés"),
    ("Compatibilité", "Flutter Web, API HTTP et base MySQL locale"),
], [2100,7020], font_size=10, caption_text="Tableau 4. Exigences non fonctionnelles.")

heading("3.1.4 Règles métier structurantes",3)
bullets([
    "une note brouillon n'est pas visible par l'étudiant et une note publiée est verrouillée ;",
    "un enseignant ne modifie que les évaluations ou publications dont il est l'auteur dans un cours attribué ;",
    "un étudiant ne consulte que ses propres notes, présences, réclamations, enrôlements et projets ;",
    "les crédits d'un cours sont acquis lorsque le seuil réglementaire est atteint ;",
    "l'appariteur publie une délibération déjà clôturée mais ne choisit pas la décision du jury ;",
    "un enrôlement valide relie un étudiant, une promotion et une année, avec une référence unique ;",
    "un projet actif exige un enrôlement valide et l'encadreur doit être compatible avec le type du projet ;",
    "les actions sensibles tirent l'identité et le rôle du jeton authentifié, jamais d'un identifiant d'autorité fourni par Flutter."
])

heading("3.2 Architecture de la solution",2)
figure(ARCH,"Figure 3. Architecture logique de Smart Faculty. Source : élaboration à partir des documents de conception.",15.3,"Architecture à trois niveaux : Flutter, FastAPI et MySQL")
heading("3.2.1 Couche de présentation",3)
p("Le frontend Flutter est organisé autour de la configuration globale, des modèles, services, fonctionnalités, composants partagés et routes. Les écrans sont regroupés par domaine et adaptés au rôle actif. Les services encapsulent les appels HTTP, la classification des erreurs et la restauration de session. L'interface ne recalcule pas les décisions métier et ne transmet pas d'identifiant libre comme autorité.")

heading("3.2.2 Couche applicative",3)
p("Le backend FastAPI constitue la source d'autorité. Les routes reçoivent la requête, les schémas valident les données, les dépendances vérifient l'identité, les services appliquent les règles et les dépôts interrogent la base. Les modules actifs couvrent notamment l'authentification, les inscriptions, les enseignants, la Valve, les notes, les résultats, les délibérations, les enrôlements et les projets.")

heading("3.2.3 Couche de données",3)
p("MySQL stocke les entités relationnelles. Le projet utilise SQLAlchemy pour les modèles et Alembic pour les migrations. InnoDB est imposé dans l'environnement de test afin de permettre les transactions et retours arrière. Les sauvegardes précèdent les migrations de la base principale ; les migrations nouvelles sont d'abord vérifiées sur smart_faculty_test par montée, descente et remontée de version.")
table(["Domaine", "Entités représentatives", "Règle d'intégrité"], [
    ("Identité", "utilisateurs, rôles, demandes d'inscription", "Adresse électronique et référence uniques"),
    ("Structure", "années, promotions, cours, affectations", "Cohérence de la période et des relations"),
    ("Évaluation", "évaluations, notes, résultats", "Pondérations, verrouillage et inscriptions actives"),
    ("Communication", "publications, pièces jointes, notifications", "Cours et auteur obligatoires"),
    ("Délibération", "sessions, membres, décisions, snapshots", "Version officielle immuable après clôture"),
    ("Enrôlement", "enrôlements académiques", "Un seul enrôlement non annulé par triplet"),
    ("Projet", "projets, encadrements, spécialités", "Un principal actif et compatibilité du domaine"),
], [1700,4050,3370], font_size=9.5, caption_text="Tableau 5. Domaines de données principaux.")

heading("3.2.4 Sécurité transversale",3)
p("La sécurité repose sur plusieurs couches complémentaires. Les mots de passe sont hachés ; les jetons d'accès et de rafraîchissement sont signés, expirables et révocables ; le compte doit rester actif ; le rôle est vérifié à chaque requête ; les routes filtrent les données par l'identité issue du jeton. Les réponses excluent les mots de passe, hachages et secrets. La politique CORS autorise les origines locales de développement selon une expression contrôlée et conserve une liste explicite en production.")

heading("3.3 Étapes de mise en œuvre",2)
heading("3.3.1 Audit et clarification de l'existant",3)
p("Le premier audit a révélé plusieurs points d'entrée backend, des technologies historiques, des scripts contradictoires, une base de test non préparée et des artefacts locaux. La réorganisation a établi frontend, backend, docs, legacy et scripts comme répertoires de référence. FastAPI a été confirmé comme backend actif ; PHP et Flask ont été archivés. Cette étape a réduit le risque de développer ou de tester le mauvais composant.")

heading("3.3.2 Stabilisation de l'environnement de test",3)
p("La base smart_faculty_test a été créée séparément de la base principale. Une protection interdit toute cible dont le nom ne se termine pas par _test. Les tests qui persistaient des données ont été isolés par transactions, savepoints et retour arrière final. Le passage des tables de test de MyISAM à InnoDB a rendu le nettoyage transactionnel effectif. Trois exécutions consécutives des vingt-six tests historiques ont alors produit le même résultat.")

heading("3.3.3 Authentification, rôles et inscriptions",3)
p("Le système de connexion a été aligné sur les rôles fonctionnels et les statuts réels de la base. Flutter utilise uniquement le rôle actif retourné par FastAPI. Les demandes publiques d'inscription sont limitées aux étudiants et enseignants ; elles ne créent ni session ni rôle privilégié. Leur approbation crée le compte, le profil et le rôle au sein d'une transaction. La persistance de session a été rendue injectable et testable, puis la restauration a été conditionnée à la vérification de /auth/moi.")

heading("3.3.4 Interface et communication web",3)
p("Une palette beige et marron a été centralisée dans le thème Flutter afin d'améliorer la cohérence visuelle. Le diagnostic de communication web a ensuite distingué les véritables erreurs CORS, les statuts HTTP, les délais et l'indisponibilité du serveur. Le backend autorise les origines locales contrôlées en développement et refuse les origines externes. Les scénarios de connexion valide, mot de passe incorrect et rôle non possédé ont été vérifiés par de vrais échanges HTTP.")

heading("3.3.5 Espace enseignant, Valve et évaluations",3)
p("L'espace enseignant a été recentré sur les cours réellement affectés. La Valve permet de créer un brouillon, publier, modifier ou archiver une publication, avec des mutations réservées à son auteur. Le module d'évaluation gère les types, pondérations, listes d'étudiants, notes et verrouillages. Une note zéro reste distincte d'une note absente ; la publication d'un cours est bloquée si la pondération n'atteint pas cent pour cent ou si une note obligatoire manque.")

heading("3.3.6 Résultats et délibérations LMD",3)
p("La consolidation semestrielle a d'abord été construite comme un aperçu provisoire, sans inventer de décision officielle. Après formalisation des règles LMD, le calcul a été pondéré par les crédits et les décisions de jury ont été implémentées. La session de délibération associe les membres, désigne un président, enregistre les décisions, clôture la session et génère un snapshot officiel. Toute correction après clôture crée une nouvelle version motivée qui conserve l'ancienne.")
figure(LMD,"Figure 4. Chaîne de décision académique. Source : synthèse des règles LMD et du module de délibération.",15.3,"Workflow enseignant, moteur, jury, appariteur et étudiant")

heading("3.3.7 Enrôlements et encadrement des projets",3)
p("L'enrôlement académique a été distingué de la demande de création de compte et de l'inscription pédagogique à un cours. Il rattache officiellement l'étudiant à une promotion et une année, avec les statuts en attente, validé ou annulé. Une fiche PDF est générée en mémoire pour l'étudiant authentifié lorsque l'enrôlement est valide.")
p("Les projets académiques sont créés par l'appariteur pour un étudiant correctement enrôlé. Les types sont limités à réseaux, systèmes embarqués, intelligence artificielle et génie logiciel. Les spécialités des enseignants sont déclarées explicitement ; elles déterminent leur compatibilité avec le projet. Le modèle autorise un encadreur principal et plusieurs co-encadreurs, tout en conservant l'historique des remplacements et désactivations.")

heading("3.3.8 Chronologie synthétique",3)
table(["Étape", "Réalisations majeures", "Validation marquante"], [
    ("Audit", "Cartographie, documentation et réorganisation", "26 tests backend stabilisés"),
    ("Authentification", "Rôles, statuts, inscription et session", "57 tests backend, 15 Flutter"),
    ("Web et interface", "Thème, CORS et erreurs réseau", "61 tests backend, 24 Flutter"),
    ("Enseignant", "Cours, Valve, évaluations et résultats", "73 tests backend, 35 Flutter"),
    ("LMD", "Consolidation, jury, snapshot et publication", "107 tests backend, 37 Flutter"),
    ("Projets", "Encadrements enseignants", "120 tests backend, 39 Flutter"),
    ("Enrôlements", "Gestion appariteur", "128 tests backend, 42 Flutter"),
    ("Encadrements", "Spécialités et attribution", "134 tests backend, 44 Flutter"),
    ("Espace étudiant", "Fiche, projet et encadreurs", "141 tests backend, 47 Flutter"),
], [1550,4700,2870], font_size=9.5, caption_text="Tableau 6. Principales étapes de réalisation.")

heading("3.4 Résultats obtenus",2)
heading("3.4.1 Résultats fonctionnels",3)
p("La version documentée de Smart Faculty offre un socle cohérent couvrant l'identité, la communication pédagogique, l'évaluation, la consolidation académique, la délibération, l'enrôlement et l'encadrement de projets. Les parcours sont filtrés par rôle et disposent d'états explicites en cas de chargement, absence de données, erreur, accès refusé ou session expirée.")
bullets([
    "authentification et restauration de session avec contrôle du rôle actif ;",
    "demandes d'inscription et approbation sécurisée ;",
    "espace enseignant fondé sur les cours réellement affectés ;",
    "Valve avec brouillons, publication et contrôle de l'auteur ;",
    "évaluations, saisie des notes, calcul, publication et verrouillage ;",
    "aperçus semestriels et décisions de jury versionnées ;",
    "gestion et consultation des enrôlements avec fiche PDF ;",
    "gestion des projets, spécialités et encadrements ;",
    "consultation étudiante de son projet et de ses encadreurs actifs."
])

heading("3.4.2 Résultats de validation",3)
figure(TESTS,"Figure 5. Progression des suites automatisées. Source : journal de développement, état au 14 juillet 2026.",15.3,"Courbes de progression des tests backend et Flutter")
table(["Contrôle", "Dernier résultat documenté", "Interprétation"], [
    ("Tests backend", "141 réussis, deux exécutions", "Services, autorisations et règles couverts"),
    ("Tests Flutter", "47 réussis, deux exécutions", "Services et comportements d'interface stables"),
    ("Analyse statique", "0 erreur, 0 avertissement", "Informations historiques non bloquantes seulement"),
    ("Build Web", "Compilation release réussie", "Application distribuable comme artefact web"),
    ("Santé FastAPI", "Trois endpoints en HTTP 200", "Application et base accessibles"),
    ("Migration", "Cycles downgrade / upgrade validés", "Évolutions de schéma réversibles en test"),
    ("Fiche PDF", "2 pages sans chevauchement", "Génération en mémoire visuellement vérifiée"),
], [1900,3000,4220], font_size=10, caption_text="Tableau 7. Résultats de validation.")
p("Ces résultats ne signifient pas que le produit est achevé dans tous ses scénarios futurs. Ils démontrent que le périmètre réalisé est soutenu par des vérifications répétables et que les limitations sont explicitement documentées.")

heading("4. Difficultés rencontrées et limites",1)
heading("4.1 Difficultés rencontrées",2)
p("Le projet a rencontré des difficultés techniques et méthodologiques significatives. Leur traitement a renforcé la compréhension de l'architecture et la discipline de validation.")
table(["Difficulté", "Réponse apportée", "Enseignement"], [
    ("Coexistence de FastAPI, Flask et PHP", "Choix officiel de FastAPI et archivage du code historique", "Clarifier la source active avant de développer"),
    ("Tests lancés sur une base inadéquate", "Protection du suffixe _test et script de préparation", "La sécurité des données commence dans l'outillage"),
    ("Données persistantes entre tests", "InnoDB, transaction externe, savepoints et rollback", "Un test doit être indépendant et reproductible"),
    ("Verrous du SDK Flutter", "Diagnostic ciblé et nettoyage des verrous orphelins", "Distinguer problème d'environnement et défaut applicatif"),
    ("Blocage CORS masqué par le navigateur", "Politique locale contrôlée et erreurs réseau typées", "Observer le preflight et les statuts réels"),
    ("Règles LMD initialement incomplètes", "Arrêt de l'implémentation puis formalisation documentaire", "Ne pas inventer une règle académique sensible"),
    ("Migrations sur tables historiques", "Sauvegarde et conversion ciblée vers InnoDB", "Préparer la compatibilité avant les clés étrangères"),
    ("Autorité transmise par le client", "Identité et rôle systématiquement dérivés du jeton", "Le backend doit rester la source d'autorité"),
], [2400,3900,2820], font_size=9.2, caption_text="Tableau 8. Difficultés, réponses et enseignements.")

heading("4.2 Limites de la solution",2)
table(["Limite actuelle", "Conséquence", "Évolution envisagée"], [
    ("Périmètre limité à une faculté", "Pas de gestion multi-facultés complète", "Paramétrage institutionnel et multi-tenant"),
    ("Déploiement surtout local / web", "Disponibilité et charge réelles non démontrées", "Hébergement, supervision et tests de performance"),
    ("Modèle LMD simplifié", "Le cours représente provisoirement une unité d'enseignement", "Décomposition élément constitutif - unité - bloc"),
    ("Progression annuelle et seconde session incomplètes", "Décisions futures non couvertes", "Moteur annuel, rattrapage et seconde session"),
    ("Présences et risques encore partiels", "Accompagnement prédictif limité", "Seuils configurables et analyses plus riches"),
    ("Projet sans collaboration avancée", "Pas de fichiers, réunions, messagerie ou notation", "Espace de suivi complet des livrables"),
    ("Modules hors périmètre", "Paiements, bibliothèque et enseignement à distance absents", "Intégrations par étapes selon les priorités"),
    ("Sécurité de production non auditée", "Aucune certification ni test d'intrusion", "Audit de sécurité, secrets gérés et journal centralisé"),
    ("Peu de preuves visuelles dans le corpus", "La monographie ne présente pas de captures applicatives", "Captures annotées et tests utilisateurs lors de la soutenance"),
], [2500,3000,3620], font_size=9.1, caption_text="Tableau 9. Limites actuelles et évolutions associées.")
p("La reconnaissance de ces limites ne réduit pas la valeur du travail. Elle définit au contraire une frontière honnête entre ce qui est validé, ce qui est prototypé et ce qui relève d'une version future.")

heading("5. Conclusion générale",1)
heading("5.1 Récapitulatif",2)
p("Cette monographie a présenté la conception et la réalisation de Smart Faculty, une plateforme destinée à moderniser la gestion académique d'une faculté universitaire. Le travail est parti d'un besoin de centralisation et de traçabilité, puis a structuré les rôles, les processus et les données autour d'une architecture client-serveur. Flutter assure l'interface, FastAPI applique les règles et MySQL conserve les informations relationnelles.")
p("La démarche a accordé une place centrale à la sécurité et à la preuve. Les identités et responsabilités sont vérifiées côté backend, les migrations sont versionnées, la base de test est isolée et les suites automatisées accompagnent chaque module. Les fonctionnalités réalisées couvrent l'authentification, l'inscription, la Valve, les évaluations, les résultats, les délibérations, les enrôlements, les projets et la consultation étudiante.")

heading("5.2 Apports du projet",2)
p("Sur le plan technique, le projet a permis de mobiliser la conception d'API, la modélisation relationnelle, les migrations, l'authentification par jetons, la gestion de rôles, les tests et la génération documentaire. Sur le plan méthodologique, il a développé la capacité à auditer un existant, préserver les données, formaliser une règle avant de la coder, documenter une décision et distinguer un résultat provisoire d'une décision officielle.")
p("Pour la faculté, Smart Faculty fournit une base intégrée susceptible de réduire les tâches manuelles, améliorer la transparence et accélérer l'accès à l'information. Pour l'étudiant, la plateforme améliore la visibilité sur son parcours. Pour l'enseignant et les responsables, elle clarifie les responsabilités et fournit des opérations contrôlées.")

heading("5.3 Perspectives",2)
bullets([
    "déployer la plateforme dans un environnement sécurisé avec sauvegardes, supervision et mesure de performance ;",
    "compléter la modélisation LMD par les éléments constitutifs, unités d'enseignement et blocs de compétences ;",
    "intégrer la progression annuelle, la seconde session et les relevés officiels ;",
    "enrichir les présences, les alertes et Campus Analytics avec des indicateurs explicables ;",
    "ajouter le dépôt de livrables, la messagerie et le suivi des réunions de projet ;",
    "réaliser des tests utilisateurs auprès des étudiants, enseignants et appariteurs ;",
    "préparer une version mobile distribuable et une stratégie de fonctionnement en connectivité limitée ;",
    "envisager un assistant académique après consolidation des règles, de la qualité des données et de la gouvernance."
])
callout("BILAN", "Smart Faculty constitue un socle professionnel, cohérent et extensible. Sa principale force est l'alignement entre règles documentées, autorisations backend, structure des données et preuves de validation.", PALE_BLUE, NAVY)

heading("6. Bibliographie et webographie",1)
heading("6.1 Bibliographie et documents de projet",2)
refs=[
    "Équipe Smart Faculty. Présentation, vision, périmètre, utilisateurs et planification du projet. Documents internes, 2026.",
    "Équipe Smart Faculty. Analyse fonctionnelle, besoins, cas d'utilisation, scénarios et règles métier. Documents internes, 2026.",
    "Équipe Smart Faculty. Architecture générale, base de données, organisation du projet et API REST. Documents internes, 2026.",
    "Équipe Smart Faculty. Cahier des charges technique de Smart Faculty. Document interne évolutif, version du 14 juillet 2026.",
    "Équipe Smart Faculty. Journal de développement de Smart Faculty. Document interne évolutif, 10-14 juillet 2026.",
    "Équipe Smart Faculty. Règles académiques LMD appliquées par Smart Faculty. Document interne, 2026.",
    "Équipe Smart Faculty. Enrôlements académiques - MVP Appariteur. Document interne, 2026.",
    "Équipe Smart Faculty. Projets et encadrements - MVP enseignant et appariteur. Document interne, 2026.",
    "Faculté des Sciences Informatiques, Université Protestante au Congo. Guide de rédaction de monographie, Licences 1 et 2 - LMD. Année académique 2024-2025.",
]
for ref in refs:
    q=p(ref); q.paragraph_format.left_indent=Cm(0.75); q.paragraph_format.first_line_indent=Cm(-0.75); q.paragraph_format.space_after=Pt(8)

heading("6.2 Textes réglementaires",2)
refs=[
    "République démocratique du Congo. Décret numéro 22/39 du 8 décembre 2022 portant organisation et fonctionnement du système LMD.",
    "Ministère de l'Enseignement supérieur et universitaire. Arrêté ministériel numéro 093/MINESU/CAB.MIN/MNB/RMM/2023 du 10 février 2023 portant cadre normatif du système LMD.",
    "Ministère de l'Enseignement supérieur et universitaire. Arrêté ministériel numéro 401/MINESU/CABMIN/MNB/RMM/MKK/2023 du 28 août 2023 portant modalités d'évaluation, de progression et d'orientation en Licence et Maîtrise.",
    "Ministère de l'Enseignement supérieur, universitaire, recherche scientifique et innovations. Instruction académique numéro 027 pour l'année académique 2025-2026, octobre 2025.",
]
for ref in refs:
    q=p(ref); q.paragraph_format.left_indent=Cm(0.75); q.paragraph_format.first_line_indent=Cm(-0.75); q.paragraph_format.space_after=Pt(8)

heading("6.3 Webographie",2)
webrefs=[
    "Flutter. Architecting Flutter apps. https://docs.flutter.dev/app-architecture, consulté le 14 juillet 2026.",
    "Flutter. Guide to app architecture. https://docs.flutter.dev/app-architecture/guide, consulté le 14 juillet 2026.",
    "FastAPI. Features. https://fastapi.tiangolo.com/features/, consulté le 14 juillet 2026.",
    "Oracle. MySQL 8.4 Reference Manual. https://dev.mysql.com/doc/refman/8.4/en/, consulté le 14 juillet 2026.",
    "Ministère de l'Enseignement supérieur, universitaire, recherche scientifique et innovations. Instruction académique numéro 027. https://minesursi.gouv.cd/images/INSTRUCTION%20027.pdf, consulté le 14 juillet 2026.",
]
for ref in webrefs:
    q=p(ref); q.paragraph_format.left_indent=Cm(0.75); q.paragraph_format.first_line_indent=Cm(-0.75); q.paragraph_format.space_after=Pt(8)

heading("7. Annexes",1)
heading("Annexe A - Matrice de traçabilité du corpus",2)
p("Cette annexe montre comment les documents présents dans le dossier docs ont été pris en compte dans la monographie.")
source_rows=[
    ("00.01 - Présentation du projet", "Contexte, problématique, solution et public"),
    ("00.02 - Vision et objectifs", "Vision, objectif général et objectifs spécifiques"),
    ("00.03 - Périmètre", "Fonctions incluses, exclusions et évolutions"),
    ("00.04 - Organigramme", "Rôles et responsabilités des utilisateurs"),
    ("00.06 - Planification", "Phases de réalisation"),
    ("01.01 - Analyse fonctionnelle", "Modules et interactions"),
    ("01.02 - Besoins fonctionnels", "Services attendus"),
    ("01.03 - Besoins non fonctionnels", "Qualité, sécurité et maintenabilité"),
    ("01.04 - Cas d'utilisation", "Actions autorisées par acteur"),
    ("01.06 - Scénarios", "Parcours principaux"),
    ("01.07 - Règles métier", "Contraintes et autorisations"),
    ("02.01 - Architecture générale", "Séparation Flutter, FastAPI et MySQL"),
    ("02.02 - Base de données", "Entités, relations et intégrité"),
    ("02.03 - Architecture du projet", "Organisation des composants"),
    ("02.04 - API REST", "Modules de routes et format d'échange"),
    ("Cahier des charges technique", "Décisions, contraintes et état technique"),
    ("Journal de développement", "Chronologie, difficultés et preuves de validation"),
    ("Règles LMD RDC", "Calcul, crédits, jury et publication"),
    ("Enrôlements académiques MVP", "Workflow appariteur et consultation étudiante"),
    ("Projets et encadrements MVP", "Types, spécialités, attribution et consultation"),
    ("Architecture frontend", "Organisation Flutter et charte initiale"),
    ("Modélisation Merise", "Domaines de données et cardinalités"),
]
table(["Source", "Apport à la monographie"],source_rows,[3700,5420],font_size=9.2)

heading("Annexe B - Routes représentatives",2)
table(["Domaine", "Route représentative", "Finalité"], [
    ("Authentification", "POST /api/v1/auth/connexion", "Ouvrir une session avec un rôle vérifié"),
    ("Inscription", "POST /api/v1/inscriptions/demandes", "Créer une demande publique"),
    ("Enseignant", "GET /api/v1/enseignants/moi", "Consulter le profil issu du jeton"),
    ("Valve", "POST /api/v1/enseignant/valve", "Créer une publication ou un brouillon"),
    ("Résultats", "POST /enseignant/cours/{id}/resultats/publier", "Publier et verrouiller un cours"),
    ("Délibération", "Routes sous /api/v1/deliberations", "Organiser jury, clôture et publication"),
    ("Enrôlement", "GET /api/v1/etudiants/moi/enrolements", "Consulter son historique"),
    ("Fiche", "GET /api/v1/etudiants/moi/enrolements/{id}/fiche", "Télécharger la fiche validée"),
    ("Projet", "GET /api/v1/etudiants/moi/projets", "Consulter ses projets non archivés"),
    ("Encadrement", "GET /api/v1/enseignants/moi/encadrements", "Consulter les projets attribués"),
], [1800,4300,3020], font_size=9.2)

heading("Annexe C - Critères de recette",2)
bullets([
    "une tentative d'accès avec un rôle non possédé est refusée par le backend ;",
    "aucun mot de passe, hachage ou jeton n'apparaît dans une réponse fonctionnelle ;",
    "un enseignant ne modifie pas la publication ou l'évaluation d'un collègue ;",
    "une note absente bloque la publication tandis qu'une note zéro reste valide ;",
    "un étudiant ne consulte jamais les données d'un autre étudiant ;",
    "une délibération publiée correspond à une session clôturée et à un snapshot versionné ;",
    "un enrôlement annulé reste historisé et ne produit pas de fiche officielle ;",
    "un projet exige un enrôlement valide et un encadreur compatible ;",
    "les tests automatisés ciblent uniquement smart_faculty_test ;",
    "toute migration majeure peut être testée par descente et remontée de version."
])

heading("Annexe D - Informations à compléter avant dépôt",2)
callout("COUVERTURE", "Renseigner le nom complet de l'étudiant, le matricule, l'encadreur et sa qualité.", PALE_GOLD, GOLD)
bullets([
    "confirmer l'intitulé exact de la filière et de l'option ;",
    "faire valider le titre final par l'encadreur ;",
    "insérer, si disponibles, des captures d'écran annotées de l'application réelle ;",
    "mettre à jour automatiquement la table des matières dans Microsoft Word ;",
    "relire les références réglementaires avec l'encadreur avant impression ;",
    "appliquer les éventuelles consignes supplémentaires de la faculté qui ne figurent pas dans le guide fourni."
])

# Prevent heading 1 on the very first body chapter from inheriting a blank-page artifact.
settings = doc.settings._element
upd = settings.find(qn("w:updateFields"))
if upd is None:
    upd=OxmlElement("w:updateFields"); settings.append(upd)
upd.set(qn("w:val"),"true")

# Ensure first body heading remains at section start while all later H1s start new pages.
first_body_h1 = next(par for par in doc.paragraphs if par.text == "1. Introduction générale")
pPr=first_body_h1._p.get_or_add_pPr()
pageBreak=pPr.find(qn("w:pageBreakBefore"))
if pageBreak is None:
    pageBreak=OxmlElement("w:pageBreakBefore"); pPr.append(pageBreak)
pageBreak.set(qn("w:val"),"0")

doc.save(FINAL)
print(FINAL)
