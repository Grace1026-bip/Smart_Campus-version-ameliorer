from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document
from pypdf import PdfReader


def read_docx(path: Path) -> dict:
    doc = Document(path)
    blocks: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append(text)
    tables: list[list[list[str]]] = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append(rows)
    return {
        "paragraphs": blocks,
        "tables": tables,
        "inline_shapes": len(doc.inline_shapes),
        "sections": len(doc.sections),
    }


def main(root: Path, guide: Path, out: Path) -> None:
    corpus = {"documents": {}, "markdown": {}, "guide": {}}
    for path in sorted(root.rglob("*.docx")):
        corpus["documents"][str(path.relative_to(root))] = read_docx(path)
    for path in sorted(root.rglob("*.md")):
        if "tmp" in path.parts or "output" in path.parts:
            continue
        corpus["markdown"][str(path.relative_to(root))] = path.read_text(encoding="utf-8")

    reader = PdfReader(guide)
    pages = []
    for index, page in enumerate(reader.pages, 1):
        pages.append({"page": index, "text": page.extract_text() or ""})
    corpus["guide"] = {"path": str(guide), "page_count": len(reader.pages), "pages": pages}
    out.write_text(json.dumps(corpus, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"Wrote {out}")
    print(f"DOCX: {len(corpus['documents'])}; Markdown: {len(corpus['markdown'])}; guide pages: {len(pages)}")


if __name__ == "__main__":
    main(Path(sys.argv[1]), Path(sys.argv[2]), Path(sys.argv[3]))
