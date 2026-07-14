from pathlib import Path
import json
from docx import Document
from pypdf import PdfReader

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "tmp" / "extracted"
OUT.mkdir(parents=True, exist_ok=True)

pdf = ROOT / "tmp" / "guide_monographie.pdf"
reader = PdfReader(str(pdf))
pages = []
for i, page in enumerate(reader.pages, 1):
    text = page.extract_text() or ""
    pages.append(f"\n===== PAGE {i} =====\n{text}")
(OUT / "guide.txt").write_text("\n".join(pages), encoding="utf-8")

index = []
for path in sorted(ROOT.rglob("*")):
    if "tmp" in path.parts or "output" in path.parts or ".git" in path.parts:
        continue
    if path.suffix.lower() == ".md":
        text = path.read_text(encoding="utf-8", errors="replace")
    elif path.suffix.lower() == ".docx":
        doc = Document(str(path))
        blocks = []
        for p in doc.paragraphs:
            if p.text.strip():
                blocks.append(p.text)
        for ti, table in enumerate(doc.tables, 1):
            blocks.append(f"\n[TABLE {ti}]")
            for row in table.rows:
                blocks.append(" | ".join(cell.text.replace("\n", " / ") for cell in row.cells))
        text = "\n".join(blocks)
    else:
        continue
    rel = path.relative_to(ROOT)
    safe = "__".join(rel.parts) + ".txt"
    (OUT / safe).write_text(text, encoding="utf-8")
    index.append({"path": str(rel), "chars": len(text), "output": safe})
(OUT / "index.json").write_text(json.dumps(index, ensure_ascii=False, indent=2), encoding="utf-8")
print(json.dumps({"pdf_pages": len(reader.pages), "sources": index}, ensure_ascii=False, indent=2))
